"""
Generate a professional .docx proposal document for the
AI-Powered Intern Progress Management System.
"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ── Color Palette ──────────────────────────────────────────────
BLUE_DARK    = RGBColor(0x1B, 0x3A, 0x5C)   # Deep navy for H1
BLUE_MED     = RGBColor(0x24, 0x4F, 0x7A)   # Medium blue for H2
BLUE_ACCENT  = RGBColor(0x2E, 0x6E, 0xA6)   # Accent blue for H3/H4
CHARCOAL     = RGBColor(0x2D, 0x2D, 0x2D)   # Near-black for body text
CHARCOAL_LT  = RGBColor(0x44, 0x44, 0x44)   # Lighter charcoal for secondary text
TABLE_HEADER = RGBColor(0x1B, 0x3A, 0x5C)   # Table header bg
TABLE_HDR_TXT= RGBColor(0xFF, 0xFF, 0xFF)   # White text on table headers
TABLE_ALT_BG = "EAF0F7"                      # Alternating row bg (hex string)
BLUE_LINE    = "1B3A5C"                       # Horizontal rule color

# ── Helpers ────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    """Set background shading on a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_table_row(table, cells_text, is_header=False, alt_row=False):
    """Add a row to a table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)

        if is_header:
            run.bold = True
            run.font.color.rgb = TABLE_HDR_TXT
            set_cell_shading(cell, BLUE_LINE)
        else:
            run.font.color.rgb = CHARCOAL
            if alt_row:
                set_cell_shading(cell, TABLE_ALT_BG)


def styled_table(doc, headers, rows):
    """Create a clean professional table."""
    table = doc.add_table(rows=0, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    # Header row
    add_table_row(table, headers, is_header=True)

    # Data rows
    for idx, row_data in enumerate(rows):
        add_table_row(table, row_data, alt_row=(idx % 2 == 1))

    # Apply borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing after table
    return table


def add_horizontal_rule(doc):
    """Add a thin blue horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{BLUE_LINE}"/>'
        '</w:pBdr>'
    )
    pPr.append(pBdr)


# ── Paragraph helpers ──────────────────────────────────────────

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = BLUE_DARK
    run.font.name = "Calibri"
    add_horizontal_rule(doc)
    return p


def heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE_MED
    run.font.name = "Calibri"
    return p


def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = BLUE_ACCENT
    run.font.name = "Calibri"
    return p


def heading4(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.color.rgb = BLUE_ACCENT
    run.font.name = "Calibri"
    return p


def para(doc, text, bold=False, italic=False, indent=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.color.rgb = CHARCOAL
    run.font.name = "Calibri"
    run.bold = bold
    run.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    indent_val = 0.35 + (level * 0.35)
    p.paragraph_format.left_indent = Inches(indent_val)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    prefix = "• " if level == 0 else "– "
    run = p.add_run(prefix + text)
    run.font.size = Pt(11)
    run.font.color.rgb = CHARCOAL
    run.font.name = "Calibri"
    return p


def bullet_bold_value(doc, label, value, level=0):
    """Bullet with bold label and normal value."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    indent_val = 0.35 + (level * 0.35)
    p.paragraph_format.left_indent = Inches(indent_val)
    p.paragraph_format.first_line_indent = Inches(-0.2)
    prefix = "• " if level == 0 else "– "
    r1 = p.add_run(prefix + label)
    r1.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = CHARCOAL
    r1.font.name = "Calibri"
    r2 = p.add_run(" " + value)
    r2.font.size = Pt(11)
    r2.font.color.rgb = CHARCOAL
    r2.font.name = "Calibri"
    return p


# ── DOCUMENT CREATION ──────────────────────────────────────────

doc = Document()

# ── Page Setup ─────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
font.color.rgb = CHARCOAL

# ══════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════

# Spacer
for _ in range(5):
    doc.add_paragraph()

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AI-Powered Intern Progress\nManagement System")
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = BLUE_DARK
run.font.name = "Calibri"

# Subtitle
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8)
run = p.add_run("Project Proposal")
run.font.size = Pt(18)
run.font.color.rgb = BLUE_MED
run.font.name = "Calibri"

# Decorative line
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
run = p.add_run("━" * 40)
run.font.size = Pt(12)
run.font.color.rgb = BLUE_ACCENT

# Meta info
for _ in range(3):
    doc.add_paragraph()

meta_items = [
    ("Prepared for:", "NETSOL Technologies"),
    ("Date:", "August 2026"),
    ("Version:", "1.0"),
    ("Status:", "Draft — Awaiting Approval"),
]
for label, value in meta_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p.add_run(label + "  ")
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = CHARCOAL
    r1.font.name = "Calibri"
    r2 = p.add_run(value)
    r2.font.size = Pt(12)
    r2.font.color.rgb = CHARCOAL_LT
    r2.font.name = "Calibri"

# Page break after cover
doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (manual)
# ══════════════════════════════════════════════════════════════

heading1(doc, "Table of Contents")

toc_items = [
    "1.  Executive Summary",
    "2.  Problem Statement",
    "3.  Proposed Solution",
    "4.  Objectives",
    "5.  Scope & Boundaries",
    "6.  Technology Stack",
    "7.  User Roles & Permissions",
    "8.  System Architecture",
    "9.  Feature Breakdown",
    "    9.1  Authentication & Onboarding Flow",
    "    9.2  Admin Panel",
    "    9.3  Intern Dashboard",
    "    9.4  Mentor Dashboard",
    "    9.5  AI-Powered Insights Engine",
    "10. Interface Descriptions",
    "11. Data Models",
    "12. API Design Overview",
    "13. AI Integration Details",
    "14. Security Considerations",
    "15. Deployment Strategy",
    "16. Future Extensibility",
    "17. Timeline & Milestones",
    "18. Conclusion",
]

for item in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if item.startswith("    "):
        p.paragraph_format.left_indent = Inches(0.5)
        run = p.add_run(item.strip())
        run.font.color.rgb = CHARCOAL_LT
    else:
        run = p.add_run(item)
        run.font.color.rgb = CHARCOAL
    run.font.size = Pt(12)
    run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════

heading1(doc, "1. Executive Summary")

para(doc, (
    "This document presents the proposal for an AI-Powered Intern Progress Management System — "
    "a purpose-built platform designed to streamline how NETSOL manages, monitors, and evaluates "
    "intern performance across their internship lifecycle."
))

para(doc, (
    "The system introduces three distinct user roles — Admin, Mentor, and Intern — each with a "
    "tailored dashboard and feature set. An integrated AI layer analyzes weekly reports, identifies "
    "at-risk interns, and generates end-of-internship performance summaries, drastically reducing "
    "the manual overhead involved in intern management."
))

para(doc, (
    "The platform is designed to complement NETSOL's existing portal (which handles attendance "
    "and other HR functions) without duplicating any of its capabilities. Instead, it focuses "
    "exclusively on work tracking, progress monitoring, learning documentation, blocker resolution, "
    "and mentor feedback."
))

# ══════════════════════════════════════════════════════════════
# 2. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════

heading1(doc, "2. Problem Statement")

para(doc, "Managing interns at scale presents several recurring challenges:")

problems = [
    ("Lack of Centralized Tracking:", "Intern tasks, progress, and feedback are often scattered across spreadsheets, emails, and verbal updates, making it difficult to get a unified view."),
    ("Inconsistent Reporting:", "Without a structured reporting mechanism, weekly progress updates vary in quality, format, and completeness."),
    ("Delayed Identification of Issues:", "Mentors may not notice an intern struggling until weeks into the program, by which time recovery is difficult."),
    ("Manual Evaluation Burden:", "At the end of an internship, mentors must manually compile feedback and assessments, which is time-consuming and prone to gaps."),
    ("No Data-Driven Insights:", "There is no mechanism to automatically surface trends, patterns, or anomalies in intern performance data."),
    ("Onboarding Friction:", "Setting up intern accounts and assigning them to mentors involves manual coordination with no standardized workflow."),
]
for label, value in problems:
    bullet_bold_value(doc, label, value)

# ══════════════════════════════════════════════════════════════
# 3. PROPOSED SOLUTION
# ══════════════════════════════════════════════════════════════

heading1(doc, "3. Proposed Solution")

para(doc, (
    "The proposed system addresses every challenge listed above through a structured, "
    "role-based web application with AI capabilities:"
))

styled_table(doc,
    ["Challenge", "Solution"],
    [
        ["Scattered tracking", "Centralized dashboard per intern with projects, tasks, and weekly reports"],
        ["Inconsistent reporting", "Structured weekly report forms with guided fields (tasks, learnings, blockers)"],
        ["Delayed issue detection", "AI-powered alerts and risk indicators visible on the mentor dashboard"],
        ["Manual evaluation", "Auto-generated end-of-internship performance summaries"],
        ["No insights", "AI analysis engine that processes weekly data and surfaces actionable insights"],
        ["Onboarding friction", "Admin panel with signup approval workflow and bulk import capability"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 4. OBJECTIVES
# ══════════════════════════════════════════════════════════════

heading1(doc, "4. Objectives")

para(doc, "The primary objectives of the system are to:")

objectives = [
    ("Centralize intern progress data", "Provide a single source of truth for all internship-related work, eliminating scattered records."),
    ("Standardize weekly reporting", "Give interns a structured, guided format for submitting weekly updates covering tasks, learnings, and blockers."),
    ("Empower mentors with visibility", "Deliver a comprehensive, at-a-glance view of each intern's six-week journey, including tasks, reports, and feedback history."),
    ("Automate performance analysis", "Leverage AI to generate insights from weekly reports, identify struggling interns early, and produce end-of-internship summaries."),
    ("Streamline administration", "Provide admins with tools for account management, signup approvals, bulk imports, and system-wide performance oversight."),
    ("Reduce manual overhead", "Minimize the time mentors and administrators spend on repetitive monitoring and evaluation tasks."),
    ("Enable future integration", "Architect the system so it can be extended and integrated into NETSOL's existing infrastructure when needed."),
]
for i, (label, value) in enumerate(objectives, 1):
    bullet_bold_value(doc, f"{label} —", value)

# ══════════════════════════════════════════════════════════════
# 5. SCOPE & BOUNDARIES
# ══════════════════════════════════════════════════════════════

heading1(doc, "5. Scope & Boundaries")

heading2(doc, "5.1 In Scope")
in_scope = [
    "User authentication and role-based access control (Admin, Mentor, Intern)",
    "Admin panel for account management, approvals, and system oversight",
    "Intern profile management and internship/project information",
    "Task creation, assignment, and status tracking",
    "Structured weekly report submission (tasks completed, learnings, blockers)",
    "Mentor feedback per week",
    "AI-generated progress insights and performance summaries",
    "Internship timeline and week-by-week progress visualization",
    "End-of-internship report generation",
    "Bulk intern import functionality for admins",
]
for item in in_scope:
    bullet(doc, item)

heading2(doc, "5.2 Out of Scope")
out_scope = [
    ("Attendance tracking", "Already handled by NETSOL's existing portal"),
    ("Leave management", "Part of existing HR systems"),
    ("Payroll or stipend management", "Outside the scope of this system"),
    ("Communication/chat features", "Can be added in future phases"),
    ("Video conferencing integration", "Not included in v1.0"),
    ("Mobile application", "Web-responsive design will serve mobile users; a dedicated app can be considered for future phases"),
]
for label, note in out_scope:
    bullet_bold_value(doc, f"{label} —", note)

# ══════════════════════════════════════════════════════════════
# 6. TECHNOLOGY STACK
# ══════════════════════════════════════════════════════════════

heading1(doc, "6. Technology Stack")

styled_table(doc,
    ["Layer", "Technology", "Rationale"],
    [
        ["Frontend", "React / Next.js", "Server-side rendering, file-based routing, excellent developer experience"],
        ["Styling", "Tailwind CSS", "Utility-first CSS framework enabling rapid, consistent, and responsive UI development"],
        ["Backend", "FastAPI (Python)", "High-performance async framework with built-in OpenAPI documentation"],
        ["Database", "PostgreSQL", "Robust, open-source relational database with excellent complex query support"],
        ["ORM", "SQLAlchemy + Alembic", "Industry-standard Python ORM with migration support"],
        ["Authentication", "JWT (JSON Web Tokens)", "Stateless authentication suitable for API-first architecture"],
        ["AI Layer", "OpenAI API / LLM Integration", "For analyzing weekly reports and generating performance insights"],
        ["Deployment", "Docker + Docker Compose", "Containerized deployment for consistency across environments"],
        ["Version Control", "Git (GitHub/GitLab)", "Standard version control with branching and PR workflows"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 7. USER ROLES & PERMISSIONS
# ══════════════════════════════════════════════════════════════

heading1(doc, "7. User Roles & Permissions")

para(doc, "The system defines three distinct roles, each with progressively broader access:")

# 7.1 Intern
heading2(doc, "7.1 Intern")
para(doc, (
    "The intern role is the primary data contributor. Interns interact with the system to "
    "document their work, report progress, and track their own growth."
))
heading3(doc, "Permissions")
intern_perms = [
    "Create and update their own profile",
    "Add internship and project details",
    "Create, update, and manage their own tasks",
    "Submit weekly reports (tasks, learnings, blockers)",
    "View mentor feedback on their reports",
    "View their own progress timeline and performance data",
    "Cannot view other interns' data",
    "Cannot access admin or mentor features",
]
for item in intern_perms:
    bullet(doc, item)

# 7.2 Mentor
heading2(doc, "7.2 Mentor")
para(doc, (
    "The mentor role is the primary evaluator and supervisor. Mentors use the system to monitor "
    "their assigned interns, provide feedback, and identify issues early."
))
heading3(doc, "Permissions")
mentor_perms = [
    "View their dedicated mentor dashboard",
    "Add/assign interns using official email addresses",
    "View all interns assigned to them",
    "Access individual intern profiles, projects, and tasks",
    "View submitted weekly reports for each intern",
    "Provide written feedback for each week",
    "View AI-generated insights and risk indicators",
    "View complete internship history for any assigned intern",
    "Cannot modify intern-submitted data",
    "Cannot access admin features or other mentors' interns",
]
for item in mentor_perms:
    bullet(doc, item)

# 7.3 Admin
heading2(doc, "7.3 Admin")
para(doc, (
    "The admin role has system-wide authority. Admins manage the platform's operational aspects "
    "— from user account lifecycle to performance oversight."
))
heading3(doc, "Permissions")
admin_perms = [
    "Approve or reject signup/registration requests",
    "Create individual user accounts (intern or mentor)",
    "Perform bulk imports of intern accounts via CSV/Excel",
    "Assign or reassign interns to mentors",
    "View all users, all interns, and all mentors across the system",
    "Access system-wide performance analytics and dashboards",
    "View aggregated AI insights across all interns",
    "Manage system settings and configurations",
    "Deactivate or archive user accounts",
    "Export reports and data",
]
for item in admin_perms:
    bullet(doc, item)

# ══════════════════════════════════════════════════════════════
# 8. SYSTEM ARCHITECTURE
# ══════════════════════════════════════════════════════════════

heading1(doc, "8. System Architecture")

para(doc, (
    "The system follows a three-tier architecture with a clear separation between the frontend, "
    "backend API, and database layers."
))

heading2(doc, "8.1 Architecture Layers")

heading3(doc, "Client Layer")
para(doc, (
    "The Next.js frontend (styled with Tailwind CSS) serves three distinct dashboards — Intern, "
    "Mentor, and Admin — each tailored to the role's needs. The frontend communicates with the "
    "backend exclusively via HTTPS REST API calls."
))

heading3(doc, "API Layer")
para(doc, (
    "The FastAPI backend handles all business logic, authentication (JWT), role-based access control (RBAC), "
    "and data processing. It includes an AI service layer that interfaces with LLM APIs for report analysis "
    "and insight generation. All API endpoints are auto-documented via FastAPI's built-in OpenAPI support."
))

heading3(doc, "Data Layer")
para(doc, (
    "PostgreSQL stores all persistent data — users, profiles, projects, tasks, weekly reports, feedback, "
    "and AI-generated insights. SQLAlchemy ORM manages data access, and Alembic handles database migrations."
))

heading2(doc, "8.2 Key Architectural Principles")
principles = [
    ("API-First Design:", "The backend exposes a well-documented REST API. The frontend consumes this API exclusively."),
    ("Stateless Authentication:", "JWT-based auth ensures the backend remains stateless and horizontally scalable."),
    ("Role-Based Access Control:", "Every API endpoint enforces role checks, ensuring users can only access data appropriate to their role."),
    ("AI as a Service Layer:", "The AI engine is encapsulated as an internal service, making it easy to swap LLM providers or upgrade models."),
    ("Database Migrations:", "Alembic manages schema evolution, ensuring safe, versioned database changes."),
]
for label, value in principles:
    bullet_bold_value(doc, label, value)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 9. FEATURE BREAKDOWN
# ══════════════════════════════════════════════════════════════

heading1(doc, "9. Feature Breakdown")

para(doc, (
    "This section details every feature of the system, organized by the flow a user would experience — "
    "starting from authentication and onboarding, then branching into role-specific dashboards."
))

# ── 9.1 Authentication & Onboarding ───────────────────────────

heading2(doc, "9.1 Authentication & Onboarding Flow")

para(doc, (
    "The authentication system is the entry point for all users. It handles registration, "
    "approval, login, and initial profile setup."
))

heading3(doc, "9.1.1 Signup / Registration")
heading4(doc, "Flow")
bullet(doc, "A new user (intern or mentor) navigates to the signup page.")
bullet(doc, "They fill out a registration form with: Full Name, Official NETSOL Email Address, Password (with confirmation), Role Selection (Intern / Mentor), and Department.")
bullet(doc, 'Upon submission, the registration request is saved with a "Pending Approval" status.')
bullet(doc, 'The user sees a confirmation: "Your registration request has been submitted. You will receive an email once an admin approves your account."')
bullet(doc, "The admin receives a notification about the pending request on their dashboard.")

heading4(doc, "Validations")
bullet(doc, "Email must be a valid NETSOL domain email.")
bullet(doc, "Password must meet minimum complexity requirements (8+ characters, uppercase, lowercase, number, special character).")
bullet(doc, "Duplicate email addresses are rejected.")

heading3(doc, "9.1.2 Admin Approval Workflow")
heading4(doc, "Flow")
bullet(doc, "Admin logs into the admin panel.")
bullet(doc, "Navigates to the Pending Approvals section.")
bullet(doc, "Sees a list of pending signup requests with details (name, email, role, department, submission date).")
bullet(doc, "For each request, admin can Approve (account is activated, credentials sent via email) or Reject (with an optional reason, user is notified).")
bullet(doc, "Admin can also create accounts directly or bulk import accounts, bypassing the signup flow.")

heading3(doc, "9.1.3 Login")
heading4(doc, "Flow")
bullet(doc, "User navigates to the login page and enters email and password.")
bullet(doc, 'System validates credentials and checks account status (must be "Active").')
bullet(doc, "On successful authentication, a JWT access token and refresh token are issued.")
bullet(doc, "The user is redirected to their role-specific dashboard: Intern → Intern Dashboard, Mentor → Mentor Dashboard, Admin → Admin Panel.")
bullet(doc, "On failure, an appropriate error message is shown (invalid credentials, account pending, account rejected, account deactivated).")

heading3(doc, "9.1.4 Password Management")
bullet_bold_value(doc, "Forgot Password:", "Users can request a password reset link sent to their registered email.")
bullet_bold_value(doc, "Reset Password:", "Users follow the emailed link to set a new password.")
bullet_bold_value(doc, "Change Password:", "Logged-in users can change their password from their profile settings.")

heading3(doc, "9.1.5 First-Time Profile Setup (Interns)")
para(doc, "After an intern logs in for the first time, they are guided through a profile completion wizard:")

heading4(doc, "Step 1 — Personal Information")
bullet(doc, "Profile photo upload")
bullet(doc, "Phone number")
bullet(doc, "University / Institution name")
bullet(doc, "Degree program and current semester/year")
bullet(doc, "LinkedIn profile URL (optional)")
bullet(doc, "GitHub profile URL (optional)")

heading4(doc, "Step 2 — Internship Details")
bullet(doc, "Internship start date")
bullet(doc, "Internship end date (auto-calculated as 6 weeks from start, or manually overridden)")
bullet(doc, "Department assigned")
bullet(doc, "Mentor assigned (auto-populated if admin pre-assigned, otherwise selected from a dropdown)")

heading4(doc, "Step 3 — Project Information")
bullet(doc, "Project title")
bullet(doc, "Project description")
bullet(doc, "Technologies / tools to be used")
bullet(doc, "Project goals (what the intern aims to achieve)")

para(doc, "Once complete, the intern lands on their full dashboard.")

doc.add_page_break()

# ── 9.2 Admin Panel ───────────────────────────────────────────

heading2(doc, "9.2 Admin Panel")

para(doc, (
    "The Admin Panel is the control center for system-wide management. "
    "It is accessible only to users with the Admin role."
))

heading3(doc, "9.2.1 Admin Dashboard (Home)")
para(doc, "The admin dashboard provides a high-level overview of the entire system.")

heading4(doc, "Key Metrics (Cards/Widgets)")
bullet(doc, "Total registered interns (active / completed / archived)")
bullet(doc, "Total registered mentors")
bullet(doc, "Pending signup approvals count")
bullet(doc, "Interns currently in progress (by week: Week 1 through Week 6)")
bullet(doc, "Average performance score across all interns")
bullet(doc, 'Number of interns flagged as "needs attention" by AI')

heading4(doc, "Visual Elements")
bullet(doc, "Bar chart: Intern distribution by department")
bullet(doc, "Line chart: Signup trends over time")
bullet(doc, "Pie chart: Intern status distribution (Active, Completed, Archived)")
bullet(doc, "Recent activity feed (latest signups, approvals, report submissions)")

heading3(doc, "9.2.2 User Management")
para(doc, "This section allows the admin to manage all user accounts in the system.")

heading4(doc, "Intern Management")
bullet(doc, "Searchable, filterable, and sortable table of all interns.")
bullet(doc, "Columns: Name, Email, Department, Mentor, Internship Status, Week, Joined Date.")
bullet(doc, "Actions per intern: View full profile, Reassign to a different mentor, Deactivate account, Archive account, Reset password.")

heading4(doc, "Mentor Management")
bullet(doc, "Table of all mentors with columns: Name, Email, Department, Number of Assigned Interns, Joined Date.")
bullet(doc, "Actions per mentor: View profile, View assigned interns, Deactivate account, Reset password.")

heading4(doc, "Admin Management")
bullet(doc, "List of admin accounts.")
bullet(doc, "Ability to add new admins (restricted to existing admins only).")

heading3(doc, "9.2.3 Signup Approval Queue")
para(doc, "A dedicated page for managing pending registrations:")
bullet(doc, "List view with: Name, Email, Requested Role, Department, Date Submitted.")
bullet(doc, "Quick filters: All Pending, Interns Only, Mentors Only.")
bullet(doc, "Bulk actions: Approve Selected, Reject Selected.")
bullet(doc, "Individual actions: Approve (with optional welcome message), Reject (with required reason).")
bullet(doc, "History tab: Previously approved and rejected requests with timestamps.")

heading3(doc, "9.2.4 Bulk Import")
para(doc, "Admins can create multiple intern accounts at once:")

heading4(doc, "Process")
bullet(doc, "Download a CSV/Excel template with required columns (Name, Email, Department, Mentor Email, Start Date).")
bullet(doc, "Fill in the template with intern data.")
bullet(doc, "Upload the completed file.")
bullet(doc, "System validates the data — checks for duplicate emails, validates mentor emails, validates date formats, and reports errors row-by-row.")
bullet(doc, "On successful validation, accounts are created in bulk.")
bullet(doc, "Each intern receives an email with temporary credentials and a link to complete their profile.")

heading3(doc, "9.2.5 Performance Overview")
para(doc, "A system-wide analytics view:")
bullet_bold_value(doc, "Department-wise Performance:", "Average scores, task completion rates, and blocker frequency per department.")
bullet_bold_value(doc, "Mentor-wise Performance:", "How interns under each mentor are performing on average.")
bullet_bold_value(doc, "Cohort Analysis:", "Compare performance across different internship cohorts (batches).")
bullet_bold_value(doc, "Export:", "Download performance data as CSV or PDF reports.")

heading3(doc, "9.2.6 System Settings")
bullet(doc, "Manage department list (add/edit/remove departments)")
bullet(doc, "Configure internship duration defaults (default: 6 weeks)")
bullet(doc, "Manage email notification templates")
bullet(doc, "Configure AI settings (enable/disable AI insights, select AI model)")
bullet(doc, "View system audit logs (who did what and when)")

doc.add_page_break()

# ── 9.3 Intern Dashboard ──────────────────────────────────────

heading2(doc, "9.3 Intern Dashboard")

para(doc, (
    "The Intern Dashboard is the intern's workspace — where they manage their profile, track tasks, "
    "submit weekly reports, and view their progress. Everything is designed around the six-week internship lifecycle."
))

heading3(doc, "9.3.1 Dashboard Home")
para(doc, "The intern's landing page after login, providing an at-a-glance summary:")

heading4(doc, "Top Section — Progress Overview")
bullet_bold_value(doc, "Internship Timeline Bar:", "A visual horizontal progress bar showing all 6 weeks, with the current week highlighted. Completed weeks are marked with a green check, the current week pulses, and future weeks are grayed out.")
bullet_bold_value(doc, "Current Week Indicator:", 'Large, prominent display: "You are in Week 3 of 6."')
bullet_bold_value(doc, "Days Remaining:", "Countdown to internship end date.")

heading4(doc, "Middle Section — Quick Stats")
bullet(doc, "Tasks assigned this week vs. tasks completed")
bullet(doc, "Weekly report status: Submitted / Pending")
bullet(doc, "Unresolved blockers count")
bullet(doc, "Skills learned this week")

heading4(doc, "Bottom Section — Recent Activity")
bullet(doc, "Latest mentor feedback (if any)")
bullet(doc, "Recently updated tasks")
bullet(doc, "Upcoming deadlines")

heading3(doc, "9.3.2 Profile Management")
para(doc, "The intern can view and edit their profile at any time.")

heading4(doc, "Editable Fields")
bullet(doc, "Profile photo, Full name, Phone number")
bullet(doc, "University / Institution, Degree program and semester")
bullet(doc, "LinkedIn and GitHub URLs, Bio / About Me")

heading4(doc, "Read-Only Fields (Editable by Admin/Mentor)")
bullet(doc, "Email address (set during registration)")
bullet(doc, "Internship start and end dates")
bullet(doc, "Department and Assigned mentor")

heading3(doc, "9.3.3 Project Management")
para(doc, "Interns can manage one or more projects they are working on during their internship.")

heading4(doc, "Project Fields")
bullet(doc, "Project title and description (rich text)")
bullet(doc, "Technologies and tools used (tag-based input)")
bullet(doc, "Project status: Not Started / In Progress / Completed")
bullet(doc, "Project start and end dates")
bullet(doc, "Repository link (GitHub/GitLab URL)")
bullet(doc, "Project goals and deliverables")

heading4(doc, "Actions")
bullet(doc, "Add a new project, Edit an existing project, Mark as complete, Delete (with confirmation)")

heading3(doc, "9.3.4 Task Management")
para(doc, (
    "Tasks are the atomic units of work within a project. The task management interface allows "
    "interns to organize, track, and update their daily/weekly work."
))

heading4(doc, "Task Board View (Kanban)")
bullet(doc, "To Do — Tasks planned but not started")
bullet(doc, "In Progress — Tasks currently being worked on")
bullet(doc, "In Review — Tasks completed and awaiting mentor review")
bullet(doc, "Done — Tasks fully completed and reviewed")

heading4(doc, "Task List View (Alternative)")
bullet(doc, "Sortable table with columns: Task Title, Project, Priority, Status, Due Date, Week")
bullet(doc, "Filters: By project, by status, by priority, by week")

heading4(doc, "Task Fields")
bullet(doc, "Task title and description")
bullet(doc, "Associated project (dropdown)")
bullet(doc, "Priority: Low / Medium / High / Critical")
bullet(doc, "Status: To Do / In Progress / In Review / Done")
bullet(doc, "Assigned week (Week 1 through Week 6)")
bullet(doc, "Due date, Estimated hours, Actual hours spent")
bullet(doc, "Notes / comments")

heading3(doc, "9.3.5 Weekly Report Submission")
para(doc, (
    "Weekly reports are the core data artifact of the system. Each week, the intern submits a "
    "structured report that captures their progress, learnings, and challenges."
))

heading4(doc, "Section 1 — Tasks Completed")
bullet(doc, 'Auto-populated list of tasks marked as "Done" during this week.')
bullet(doc, "Intern can add additional context or notes to each task.")
bullet(doc, "Intern can manually add tasks not tracked in the task manager.")

heading4(doc, "Section 2 — Tasks In Progress")
bullet(doc, 'Auto-populated list of tasks currently in "In Progress" or "In Review" status.')
bullet(doc, "Expected completion timeline.")

heading4(doc, "Section 3 — Learnings & Skills Gained")
bullet(doc, "Free-text area for the intern to describe what they learned this week.")
bullet(doc, 'Tag-based skill input (e.g., "Python", "REST APIs", "Git", "Docker").')
bullet(doc, "The system maintains a cumulative skill tracker across all weeks.")

heading4(doc, "Section 4 — Blockers & Challenges")
bullet(doc, "Structured blocker entries with: description, severity (Minor/Moderate/Critical), status (Unresolved/Resolved), and help needed.")
bullet(doc, "Blockers persist across weeks until marked as resolved.")

heading4(doc, "Section 5 — Goals for Next Week")
bullet(doc, "Free-text area where the intern outlines what they plan to accomplish next week.")

heading4(doc, "Section 6 — Self-Assessment")
bullet(doc, "Rating (1–5 scale): How productive was this week?")
bullet(doc, "Rating (1–5 scale): How confident do you feel about your progress?")
bullet(doc, "Optional: Additional comments")

heading4(doc, "Submission Rules")
bullet(doc, "One report per week.")
bullet(doc, "Reports can be saved as drafts before final submission.")
bullet(doc, "Once submitted, reports cannot be edited (to preserve integrity).")
bullet(doc, "Late submissions are flagged in the system.")
bullet(doc, "If a report is not submitted by the deadline (e.g., Friday 6 PM), the mentor and admin are notified.")

heading3(doc, "9.3.6 Internship Timeline View")
para(doc, "A visual, week-by-week timeline that shows the intern's entire internship journey:")

bullet(doc, "Each week is represented as a card containing: report submission status, tasks completed, key learnings (tags), blockers, mentor feedback summary, and self-assessment scores.")
bullet(doc, "Color coding: Green (good progress), Yellow (average), Red (needs attention), Gray (future/no data).")
bullet(doc, "The current week is visually distinct and expanded by default.")

heading3(doc, "9.3.7 Mentor Feedback View")
para(doc, "Interns can view feedback their mentor has provided:")
bullet(doc, "Organized by week.")
bullet(doc, "Each feedback entry shows: week number, feedback text, rating/score, and date given.")
bullet(doc, "Read-only — interns cannot edit or respond to feedback in v1.0.")

heading3(doc, "9.3.8 Progress & Analytics")
para(doc, "A personal analytics page for the intern to visualize their own performance:")
bullet_bold_value(doc, "Task Completion Trend:", "Line chart showing tasks completed per week.")
bullet_bold_value(doc, "Skill Growth:", "Visualization of skills accumulated over the internship (tag cloud or bar chart).")
bullet_bold_value(doc, "Self-Assessment Trend:", "Line chart of weekly self-assessment scores.")
bullet_bold_value(doc, "Blocker History:", "Timeline showing when blockers were raised and resolved.")
bullet_bold_value(doc, "Overall Progress Score:", "A computed score based on task completion, report timeliness, mentor feedback, and self-assessment.")

doc.add_page_break()

# ── 9.4 Mentor Dashboard ──────────────────────────────────────

heading2(doc, "9.4 Mentor Dashboard")

para(doc, (
    "The Mentor Dashboard is designed for oversight and evaluation. Mentors use it to monitor their "
    "assigned interns, review weekly reports, provide feedback, and leverage AI insights."
))

heading3(doc, "9.4.1 Dashboard Home")
para(doc, "The mentor's landing page provides a consolidated view of all assigned interns.")

heading4(doc, "Summary Cards")
bullet(doc, "Total interns assigned")
bullet(doc, "Interns currently active (in-progress internships)")
bullet(doc, "Interns who have completed their internship")
bullet(doc, "Pending weekly reports (not yet reviewed by mentor)")
bullet(doc, 'Interns flagged as "needs attention" (by AI or by blocker severity)')

heading4(doc, "Intern Overview Table")
bullet(doc, "Searchable, sortable table listing all assigned interns.")
bullet(doc, "Columns: Intern name (clickable), profile photo, department, current week, latest report status, task completion rate, AI risk indicator (Green/Yellow/Red), last activity timestamp.")
bullet(doc, "Heatmap or sparkline per intern showing weekly performance trend.")
bullet(doc, "Alerts/badges for interns with unresolved critical blockers.")

heading3(doc, "9.4.2 Intern Detail View")
para(doc, "When a mentor clicks on an intern's name, they see a comprehensive profile page with tabbed navigation:")

heading4(doc, "Profile Header")
bullet(doc, "Intern's photo, name, email, university, department.")
bullet(doc, "Internship dates and current week.")
bullet(doc, "Overall progress bar.")

heading4(doc, "Tab 1 — Overview")
bullet(doc, "Summary statistics (total tasks, tasks completed, reports submitted, average self-assessment).")
bullet(doc, "Performance trend chart (line graph across 6 weeks).")
bullet(doc, "Cumulative skills gained (tag cloud).")

heading4(doc, "Tab 2 — Projects")
bullet(doc, "List of all projects the intern is working on with full details.")
bullet(doc, "Read-only for the mentor.")

heading4(doc, "Tab 3 — Tasks")
bullet(doc, "Full task list (all weeks), filterable by week, project, status, and priority.")
bullet(doc, "Kanban view available. Read-only for the mentor.")

heading4(doc, "Tab 4 — Weekly Reports")
bullet(doc, "Expandable accordion for each week (Week 1 through Week 6).")
bullet(doc, "Each week shows: full report content, submission date/time, on-time status, AI-generated summary, and mentor feedback input area.")

heading4(doc, "Tab 5 — Feedback History")
bullet(doc, "Chronological list of all feedback the mentor has provided for this intern.")
bullet(doc, "Editable — mentor can update feedback for any week.")

heading4(doc, "Tab 6 — AI Insights")
bullet(doc, "AI-generated analysis specific to this intern.")
bullet(doc, "Progress trajectory prediction, strength/improvement areas, risk assessment.")

heading3(doc, "9.4.3 Providing Weekly Feedback")
para(doc, "For each submitted weekly report, the mentor can provide feedback:")

heading4(doc, "Feedback Form Fields")
bullet_bold_value(doc, "Written Feedback:", "Rich text area for detailed comments, suggestions, and guidance.")
bullet_bold_value(doc, "Performance Rating:", "1–5 star scale for the week.")
bullet_bold_value(doc, "Categories (optional):", "Meeting expectations, Exceeding expectations, Needs improvement, Requires immediate attention.")
bullet_bold_value(doc, "Action Items:", "Specific things the intern should focus on next week.")

heading4(doc, "Behavior")
bullet(doc, "Feedback is saved and immediately visible to the intern.")
bullet(doc, "Mentor can edit feedback at any time.")
bullet(doc, "The system timestamps feedback submissions for audit purposes.")

heading3(doc, "9.4.4 Intern Assignment")
bullet(doc, "Mentors can add interns by entering the intern's official NETSOL email address.")
bullet(doc, "If the intern account exists and is unassigned, the assignment is made immediately.")
bullet(doc, "If assigned to another mentor, a reassignment request is sent to the admin.")
bullet(doc, "If the account does not exist, the mentor is notified and can request the admin to create it.")

heading3(doc, "9.4.5 End-of-Internship Evaluation")
para(doc, "At the end of an intern's 6-week period, the mentor is prompted to complete a final evaluation:")

heading4(doc, "Evaluation Form")
bullet(doc, "Overall Performance Rating: 1–10 scale")
bullet(doc, "Technical Skills Assessment: Rating per skill tag accumulated during the internship")
bullet(doc, "Soft Skills Assessment: Communication, initiative, teamwork, time management (1–5 each)")
bullet(doc, "Strengths and Areas for Improvement (free text)")
bullet(doc, "Recommendation: Hire / Extend Internship / Do Not Hire / Undecided")
bullet(doc, "Final Comments (free text)")

heading4(doc, "Output")
para(doc, (
    "The evaluation, combined with AI-generated insights, produces a comprehensive End-of-Internship "
    "Report downloadable as a PDF."
))

doc.add_page_break()

# ── 9.5 AI-Powered Insights Engine ────────────────────────────

heading2(doc, "9.5 AI-Powered Insights Engine")

para(doc, (
    "The AI layer is one of the system's differentiating features. It operates in the background, "
    "analyzing data submitted by interns and surfacing actionable insights for mentors and admins."
))

heading3(doc, "9.5.1 Weekly Report Analysis")
para(doc, "When an intern submits a weekly report, the AI processes it to generate:")
bullet_bold_value(doc, "Report Summary:", "A concise 2–3 sentence summary of the week's progress.")
bullet_bold_value(doc, "Sentiment Analysis:", "Detects the overall tone (positive, neutral, negative). Flags reports indicating frustration, confusion, or disengagement.")
bullet_bold_value(doc, "Blocker Severity Assessment:", "Evaluates reported blockers and flags critical ones requiring immediate mentor attention.")
bullet_bold_value(doc, "Consistency Check:", "Compares report content with task statuses to identify discrepancies.")

heading3(doc, "9.5.2 Progress Tracking & Risk Identification")
para(doc, "The AI continuously monitors intern data to identify risk patterns:")
bullet_bold_value(doc, "Declining Performance:", "Flags interns whose task completion rate, self-assessment scores, or report quality is trending downward.")
bullet_bold_value(doc, "Persistent Blockers:", "Highlights interns with blockers unresolved for multiple weeks.")
bullet_bold_value(doc, "Late or Missing Reports:", "Identifies interns who are consistently late or miss weekly submissions.")
bullet_bold_value(doc, "Low Engagement:", "Detects minimal task activity or very brief report submissions.")

heading4(doc, "Risk Levels")
styled_table(doc,
    ["Level", "Indicator", "Description"],
    [
        ["On Track", "Green", "Intern is progressing well"],
        ["Monitor", "Yellow", "Some indicators suggest the intern may need additional support"],
        ["At Risk", "Red", "Significant concerns detected; mentor should intervene"],
    ]
)

heading3(doc, "9.5.3 Comparative Insights (Mentor View)")
para(doc, "For mentors managing multiple interns, the AI provides:")
bullet_bold_value(doc, "Relative Performance:", "How each intern compares to their peers.")
bullet_bold_value(doc, "Common Blockers:", "Identifies if multiple interns face similar challenges, suggesting systemic issues.")
bullet_bold_value(doc, "Skill Gap Analysis:", "Highlights skill areas where interns are strong vs. where they need more support.")

heading3(doc, "9.5.4 End-of-Internship AI Summary")
para(doc, "At the conclusion of the internship, the AI generates a comprehensive performance report:")
bullet_bold_value(doc, "Executive Summary:", "A paragraph summarizing the intern's overall journey.")
bullet_bold_value(doc, "Week-by-Week Narrative:", "A brief summary of each week's highlights.")
bullet_bold_value(doc, "Skill Development Map:", "Visual representation of skills at start vs. end.")
bullet_bold_value(doc, "Key Achievements:", "Automatically extracted from weekly reports and task data.")
bullet_bold_value(doc, "Areas for Growth:", "Derived from blocker patterns, low-scoring weeks, and mentor feedback.")

heading4(doc, "Performance Score Breakdown")
styled_table(doc,
    ["Factor", "Weight"],
    [
        ["Task completion rate", "25%"],
        ["Report timeliness", "15%"],
        ["Blocker resolution speed", "10%"],
        ["Mentor feedback scores", "25%"],
        ["Self-assessment trend", "10%"],
        ["Skill growth", "15%"],
    ]
)

heading3(doc, "9.5.5 Admin-Level AI Analytics")
para(doc, "For admins, the AI provides system-wide insights:")
bullet_bold_value(doc, "Cohort Performance Summary:", "How the current batch of interns is performing overall.")
bullet_bold_value(doc, "Department Insights:", "Which departments have the best intern outcomes and why.")
bullet_bold_value(doc, "Mentor Effectiveness:", "Which mentors' interns show the strongest growth patterns (handled sensitively).")
bullet_bold_value(doc, "Trend Analysis:", "How intern quality and performance compare across multiple internship cycles.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 10. INTERFACE DESCRIPTIONS
# ══════════════════════════════════════════════════════════════

heading1(doc, "10. Interface Descriptions")

para(doc, "This section summarizes the key screens and their layouts across all three roles.")

heading2(doc, "10.1 Common Interfaces")

heading3(doc, "Login Page")
bullet(doc, "Clean, centered login form with Email and Password fields.")
bullet(doc, '"Remember Me" checkbox and "Forgot Password" link.')
bullet(doc, '"Sign Up" link for new users.')
bullet(doc, "NETSOL branding and logo.")

heading3(doc, "Signup Page")
bullet(doc, "Registration form with: Name, Email, Password, Confirm Password, Role (dropdown), Department (dropdown).")
bullet(doc, "Terms and conditions checkbox.")
bullet(doc, '"Already have an account? Login" link.')

heading3(doc, "Navigation Sidebar (All Roles)")
bullet(doc, "Persistent left sidebar with role-appropriate menu items.")
bullet(doc, "User avatar and name at the top.")
bullet(doc, "Collapsible for more screen space, active page indicator, logout option at the bottom.")

heading2(doc, "10.2 Admin Interfaces")
styled_table(doc,
    ["Screen", "Description"],
    [
        ["Admin Dashboard", "Metric cards, charts (department distribution, signup trends, status pie chart), activity feed"],
        ["Pending Approvals", "Table of pending signups with Approve/Reject actions, filters, bulk actions"],
        ["User Management", "Tabbed view: Interns / Mentors / Admins. Searchable tables with action buttons"],
        ["Bulk Import", "Template download, file upload area, validation results, import confirmation"],
        ["Performance Overview", "Department-wise and mentor-wise analytics with charts and data tables"],
        ["System Settings", "Form-based settings for departments, internship defaults, email templates, AI config"],
    ]
)

heading2(doc, "10.3 Intern Interfaces")
styled_table(doc,
    ["Screen", "Description"],
    [
        ["Intern Dashboard", "Timeline bar, quick stats cards, recent activity, report status"],
        ["Profile Page", "Editable form with personal and internship details"],
        ["Projects Page", "Project cards with details, add/edit/delete actions"],
        ["Task Board", "Kanban board (drag-and-drop) and list view toggle"],
        ["Weekly Report Form", "Multi-section form: tasks, learnings, blockers, goals, self-assessment"],
        ["Timeline View", "Horizontal timeline with week cards, color-coded status"],
        ["Feedback View", "Chronological feedback entries per week"],
        ["Progress Page", "Charts: task completion trend, skill growth, blocker history, overall score"],
    ]
)

heading2(doc, "10.4 Mentor Interfaces")
styled_table(doc,
    ["Screen", "Description"],
    [
        ["Mentor Dashboard", "Summary cards, intern overview table with risk indicators, alerts"],
        ["Intern Detail View", "Tabbed view: Overview / Projects / Tasks / Reports / Feedback / AI Insights"],
        ["Feedback Form", "Text area, star rating, category checkboxes, action items list"],
        ["End-of-Internship Evaluation", "Comprehensive evaluation form with ratings, assessments, and recommendation"],
        ["AI Insights Panel", "Risk indicators, trend analysis, comparative insights, skill gap analysis"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 11. DATA MODELS
# ══════════════════════════════════════════════════════════════

heading1(doc, "11. Data Models")

para(doc, "The following are the core entities that make up the system's data layer:")

heading2(doc, "11.1 Core Entities")

styled_table(doc,
    ["Entity", "Description", "Key Fields"],
    [
        ["User", "Core authentication entity", "id, email, hashed_password, role (admin/mentor/intern), status (pending/active/deactivated/archived), created_at"],
        ["Profile", "Extended user information (1:1 with User)", "user_id, phone, university, degree, semester, linkedin_url, github_url, photo_url, bio"],
        ["Internship", "An intern's internship period", "id, intern_id (FK→User), mentor_id (FK→User), department, start_date, end_date, status"],
        ["Project", "A project the intern is working on", "id, internship_id, title, description, technologies (JSON), status, repo_url, goals"],
        ["Task", "Individual unit of work within a Project", "id, project_id, intern_id, title, description, priority, status, week, due_date, estimated_hours, actual_hours"],
        ["WeeklyReport", "Intern's weekly submission", "id, internship_id, week_number, tasks_completed, tasks_in_progress, learnings, goals_next_week, self_score_productivity, self_score_confidence, submitted_at, status"],
        ["Blocker", "A challenge reported by the intern", "id, report_id, description, severity (minor/moderate/critical), status (unresolved/resolved), help_needed"],
        ["MentorFeedback", "Mentor's response to a WeeklyReport (1:1)", "id, report_id, mentor_id, feedback_text, rating, category, action_items (JSON), created_at"],
        ["AIInsight", "AI-generated analysis data", "id, report_id (nullable), intern_id, type (summary/risk/sentiment), content (JSON), risk_level, generated_at"],
        ["EndOfInternshipEvaluation", "Mentor's final assessment", "id, internship_id, mentor_id, overall_rating, technical_scores (JSON), soft_skill_scores (JSON), strengths, improvements, recommendation, ai_summary, final_score"],
    ]
)

heading2(doc, "11.2 Key Relationships")
bullet(doc, "User → Profile: One-to-One")
bullet(doc, "User (Intern) → Internship: One-to-One (an intern has one active internship)")
bullet(doc, "User (Mentor) → Internship: One-to-Many (a mentor supervises multiple interns)")
bullet(doc, "Internship → Project: One-to-Many")
bullet(doc, "Project → Task: One-to-Many")
bullet(doc, "Internship → WeeklyReport: One-to-Many (max 6 per internship)")
bullet(doc, "WeeklyReport → Blocker: One-to-Many")
bullet(doc, "WeeklyReport → MentorFeedback: One-to-One")
bullet(doc, "WeeklyReport → AIInsight: One-to-Many")
bullet(doc, "Internship → EndOfInternshipEvaluation: One-to-One")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 12. API DESIGN OVERVIEW
# ══════════════════════════════════════════════════════════════

heading1(doc, "12. API Design Overview")

para(doc, (
    "The backend API follows RESTful conventions with versioned endpoints. "
    "All endpoints require JWT authentication unless marked as public."
))

heading2(doc, "12.1 Authentication Endpoints")
styled_table(doc,
    ["Method", "Endpoint", "Description", "Access"],
    [
        ["POST", "/api/v1/auth/signup", "Register a new user", "Public"],
        ["POST", "/api/v1/auth/login", "Authenticate and receive JWT", "Public"],
        ["POST", "/api/v1/auth/refresh", "Refresh access token", "Authenticated"],
        ["POST", "/api/v1/auth/forgot-password", "Request password reset email", "Public"],
        ["POST", "/api/v1/auth/reset-password", "Reset password with token", "Public"],
        ["PUT", "/api/v1/auth/change-password", "Change current password", "Authenticated"],
    ]
)

heading2(doc, "12.2 Admin Endpoints")
styled_table(doc,
    ["Method", "Endpoint", "Description", "Access"],
    [
        ["GET", "/api/v1/admin/dashboard", "Get admin dashboard metrics", "Admin"],
        ["GET", "/api/v1/admin/approvals", "List pending signup requests", "Admin"],
        ["PUT", "/api/v1/admin/approvals/{id}/approve", "Approve a signup request", "Admin"],
        ["PUT", "/api/v1/admin/approvals/{id}/reject", "Reject a signup request", "Admin"],
        ["GET", "/api/v1/admin/users", "List all users (filterable by role)", "Admin"],
        ["POST", "/api/v1/admin/users", "Create a user account", "Admin"],
        ["PUT", "/api/v1/admin/users/{id}", "Update user details", "Admin"],
        ["DELETE", "/api/v1/admin/users/{id}", "Deactivate a user", "Admin"],
        ["POST", "/api/v1/admin/users/bulk-import", "Bulk import interns via CSV", "Admin"],
        ["GET", "/api/v1/admin/analytics", "Get system-wide analytics", "Admin"],
        ["GET", "/api/v1/admin/settings", "Get system settings", "Admin"],
        ["PUT", "/api/v1/admin/settings", "Update system settings", "Admin"],
    ]
)

heading2(doc, "12.3 Intern Endpoints")
styled_table(doc,
    ["Method", "Endpoint", "Description", "Access"],
    [
        ["GET", "/api/v1/intern/profile", "Get own profile", "Intern"],
        ["PUT", "/api/v1/intern/profile", "Update own profile", "Intern"],
        ["GET", "/api/v1/intern/internship", "Get internship details", "Intern"],
        ["GET", "/api/v1/intern/projects", "List own projects", "Intern"],
        ["POST", "/api/v1/intern/projects", "Create a project", "Intern"],
        ["PUT", "/api/v1/intern/projects/{id}", "Update a project", "Intern"],
        ["DELETE", "/api/v1/intern/projects/{id}", "Delete a project", "Intern"],
        ["GET", "/api/v1/intern/tasks", "List own tasks (filterable)", "Intern"],
        ["POST", "/api/v1/intern/tasks", "Create a task", "Intern"],
        ["PUT", "/api/v1/intern/tasks/{id}", "Update a task", "Intern"],
        ["DELETE", "/api/v1/intern/tasks/{id}", "Delete a task", "Intern"],
        ["GET", "/api/v1/intern/reports", "List own weekly reports", "Intern"],
        ["POST", "/api/v1/intern/reports", "Submit a weekly report", "Intern"],
        ["GET", "/api/v1/intern/reports/{week}", "Get report for a specific week", "Intern"],
        ["GET", "/api/v1/intern/feedback", "Get all mentor feedback", "Intern"],
        ["GET", "/api/v1/intern/progress", "Get progress analytics", "Intern"],
        ["GET", "/api/v1/intern/timeline", "Get timeline data", "Intern"],
    ]
)

heading2(doc, "12.4 Mentor Endpoints")
styled_table(doc,
    ["Method", "Endpoint", "Description", "Access"],
    [
        ["GET", "/api/v1/mentor/dashboard", "Get mentor dashboard data", "Mentor"],
        ["GET", "/api/v1/mentor/interns", "List assigned interns", "Mentor"],
        ["POST", "/api/v1/mentor/interns/assign", "Assign an intern by email", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}", "Get intern detail view", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/projects", "Get intern's projects", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/tasks", "Get intern's tasks", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/reports", "Get intern's weekly reports", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/reports/{week}", "Get specific week's report", "Mentor"],
        ["POST", "/api/v1/mentor/interns/{id}/feedback/{week}", "Submit feedback for a week", "Mentor"],
        ["PUT", "/api/v1/mentor/interns/{id}/feedback/{week}", "Update feedback for a week", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/insights", "Get AI insights for an intern", "Mentor"],
        ["POST", "/api/v1/mentor/interns/{id}/evaluation", "Submit final evaluation", "Mentor"],
        ["GET", "/api/v1/mentor/interns/{id}/report/export", "Export intern report as PDF", "Mentor"],
    ]
)

heading2(doc, "12.5 AI Endpoints")
styled_table(doc,
    ["Method", "Endpoint", "Description", "Access"],
    [
        ["POST", "/api/v1/ai/analyze-report", "Trigger AI analysis on a report", "System/Internal"],
        ["GET", "/api/v1/ai/insights/{intern_id}", "Get AI insights for an intern", "Mentor, Admin"],
        ["GET", "/api/v1/ai/summary/{intern_id}", "Get AI end-of-internship summary", "Mentor, Admin"],
        ["GET", "/api/v1/ai/cohort-analysis", "Get AI cohort-level analysis", "Admin"],
    ]
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════
# 13. AI INTEGRATION DETAILS
# ══════════════════════════════════════════════════════════════

heading1(doc, "13. AI Integration Details")

heading2(doc, "13.1 Technology")
para(doc, (
    "The AI layer integrates with a Large Language Model (LLM) via API — initially OpenAI's GPT-4 "
    "or equivalent. The integration is abstracted behind a service layer, making it straightforward "
    "to switch providers."
))

heading2(doc, "13.2 Data Pipeline")
para(doc, (
    "When an intern submits a weekly report, the system triggers the AI service. The report content "
    "is sent to the LLM API along with a carefully crafted prompt. The AI processes the report and "
    "returns structured insights (summary, sentiment, risk indicators). These results are stored in "
    "the AIInsight table and displayed on mentor and admin dashboards."
))

heading2(doc, "13.3 Prompt Engineering")
para(doc, "The system uses carefully crafted prompts for each AI function:")
bullet_bold_value(doc, "Report Summarization:", "Summarize the weekly report in 2–3 sentences, highlighting key accomplishments and concerns.")
bullet_bold_value(doc, "Sentiment Analysis:", "Analyze the tone and classify as positive, neutral, or negative. Identify language indicating frustration or disengagement.")
bullet_bold_value(doc, "Risk Assessment:", "Based on 6-week data (tasks, reports, blockers, feedback), assess performance trajectory and flag risk indicators.")
bullet_bold_value(doc, "End-of-Internship Summary:", "Generate a comprehensive performance summary based on complete 6-week data.")

heading2(doc, "13.4 Rate Limiting & Cost Management")
bullet(doc, "AI analysis is triggered only on report submission (not on every page load).")
bullet(doc, "Results are cached in the database after generation.")
bullet(doc, "AI endpoints are rate-limited to prevent abuse.")
bullet(doc, "Token usage is logged for cost tracking.")
bullet(doc, "Admins can enable/disable AI features from system settings.")

# ══════════════════════════════════════════════════════════════
# 14. SECURITY CONSIDERATIONS
# ══════════════════════════════════════════════════════════════

heading1(doc, "14. Security Considerations")

styled_table(doc,
    ["Area", "Approach"],
    [
        ["Authentication", "JWT-based with short-lived access tokens (15 min) and longer refresh tokens (7 days)"],
        ["Password Storage", "Bcrypt hashing with salt"],
        ["Authorization", "Role-based access control (RBAC) enforced at the API middleware level"],
        ["Data Isolation", "Interns see only own data. Mentors see only assigned interns. Admins have full access."],
        ["Input Validation", "All API inputs validated using Pydantic schemas (FastAPI's built-in validation)"],
        ["CORS", "Configured to allow only the frontend domain"],
        ["Rate Limiting", "Applied to auth endpoints to prevent brute force attacks"],
        ["HTTPS", "All traffic encrypted via TLS in production"],
        ["SQL Injection", "Prevented by using SQLAlchemy ORM (parameterized queries)"],
        ["XSS Protection", "React's built-in JSX escaping + Content Security Policy headers"],
        ["Audit Logging", "All admin actions and sensitive operations logged with timestamps and user IDs"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 15. DEPLOYMENT STRATEGY
# ══════════════════════════════════════════════════════════════

heading1(doc, "15. Deployment Strategy")

heading2(doc, "15.1 Environments")
styled_table(doc,
    ["Environment", "Purpose"],
    [
        ["Development", "Local development with hot-reload. Docker Compose for services."],
        ["Staging", "Pre-production testing. Mirrors production configuration."],
        ["Production", "Live environment accessible to NETSOL users."],
    ]
)

heading2(doc, "15.2 Containerization")
para(doc, "The application is fully containerized using Docker:")
bullet_bold_value(doc, "Frontend Container:", "Next.js application served via Node.js")
bullet_bold_value(doc, "Backend Container:", "FastAPI application served via Uvicorn")
bullet_bold_value(doc, "Database Container:", "PostgreSQL (or managed database service in production)")
bullet_bold_value(doc, "Reverse Proxy:", "Nginx for routing, SSL termination, and static file serving")

heading2(doc, "15.3 CI/CD Pipeline")
bullet_bold_value(doc, "Version Control:", "Git with feature branch workflow.")
bullet_bold_value(doc, "CI:", "Automated linting, testing, and build verification on every pull request.")
bullet_bold_value(doc, "CD:", "Automated deployment to staging on merge to develop, manual promotion to production from main.")

# ══════════════════════════════════════════════════════════════
# 16. FUTURE EXTENSIBILITY
# ══════════════════════════════════════════════════════════════

heading1(doc, "16. Future Extensibility")

para(doc, (
    "The system is designed with extensibility in mind. The following features are candidates for future phases:"
))

styled_table(doc,
    ["Phase", "Feature", "Description"],
    [
        ["v1.1", "Email Notifications", "Automated emails for report deadlines, feedback availability, approval status changes"],
        ["v1.1", "In-App Notifications", "Real-time notification bell with unread count for all roles"],
        ["v1.2", "Chat / Messaging", "Direct messaging between intern and mentor within the platform"],
        ["v1.2", "File Attachments", "Allow interns to attach files, screenshots, or documents to reports and tasks"],
        ["v2.0", "NETSOL Portal Integration", "SSO integration with NETSOL's existing portal, shared authentication, and attendance data sync"],
        ["v2.0", "Mobile Application", "Native mobile apps (React Native) for interns and mentors"],
        ["v2.0", "Custom Internship Durations", "Support for internships longer or shorter than 6 weeks"],
        ["v2.1", "Mentor Matching Algorithm", "AI-assisted mentor-intern matching based on skills, department, and availability"],
        ["v2.1", "Certification Generation", "Auto-generated internship completion certificates with performance grades"],
        ["v3.0", "Multi-Tenant Support", "Allow other organizations to use the platform (SaaS model)"],
        ["v3.0", "Advanced Analytics Dashboard", "Executive-level dashboards with drill-down capability and export to BI tools"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 17. TIMELINE & MILESTONES
# ══════════════════════════════════════════════════════════════

heading1(doc, "17. Timeline & Milestones")

para(doc, "The estimated development timeline for v1.0 is 8–10 weeks:")

styled_table(doc,
    ["Week", "Phase", "Deliverables"],
    [
        ["Week 1", "Project Setup & Design", "Repository setup, database schema design, UI/UX wireframes, API contract definition"],
        ["Week 2", "Authentication & Admin Core", "Signup, login, JWT auth, admin panel scaffold, approval workflow"],
        ["Week 3", "Admin Panel Completion", "User management, bulk import, settings, admin dashboard metrics"],
        ["Week 4", "Intern Dashboard — Core", "Profile management, project CRUD, task management (Kanban + list)"],
        ["Week 5", "Intern Dashboard — Reports", "Weekly report submission, timeline view, blocker tracking"],
        ["Week 6", "Mentor Dashboard — Core", "Intern overview, detail view, report review, feedback submission"],
        ["Week 7", "AI Integration", "Report analysis, risk assessment, insight generation, end-of-internship summary"],
        ["Week 8", "Analytics & Polish", "Progress charts, performance analytics, UI polish, responsive design"],
        ["Week 9", "Testing & QA", "Unit tests, integration tests, end-to-end testing, bug fixes"],
        ["Week 10", "Deployment & Launch", "Staging deployment, UAT, production deployment, documentation"],
    ]
)

# ══════════════════════════════════════════════════════════════
# 18. CONCLUSION
# ══════════════════════════════════════════════════════════════

heading1(doc, "18. Conclusion")

para(doc, (
    "The AI-Powered Intern Progress Management System is designed to transform how NETSOL manages its "
    "internship program. By providing structured workflows for interns, powerful monitoring tools for "
    "mentors, and comprehensive oversight for administrators — all enhanced by AI-driven insights — the "
    "system eliminates manual overhead, ensures no intern falls through the cracks, and produces "
    "data-driven evaluations."
))

para(doc, (
    "The chosen technology stack (Next.js, Tailwind CSS, FastAPI, PostgreSQL) ensures a modern, "
    "performant, and maintainable application. The modular architecture and API-first design make "
    "the system ready for future integration with NETSOL's existing infrastructure."
))

para(doc, (
    "This proposal serves as the foundation for development. Upon approval, the team will proceed "
    "with detailed UI/UX design, database schema finalization, and sprint planning based on the "
    "timeline outlined above."
))

# Signature area
for _ in range(3):
    doc.add_paragraph()

add_horizontal_rule(doc)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(12)
run = p.add_run("End of Proposal")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = CHARCOAL_LT
run.font.name = "Calibri"

# ══════════════════════════════════════════════════════════════
# SAVE
# ══════════════════════════════════════════════════════════════

output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), "proposal.docx")
doc.save(output_path)
print(f"✅ Document saved to: {output_path}")
